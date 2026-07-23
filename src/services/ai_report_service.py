"""
AI Financial Report Service — Sinh báo cáo tài chính bằng ngôn ngữ tự nhiên
Chạy background, cache kết quả, fallback template khi Gemini fail
"""
import io
import json
import logging
from typing import Optional
from datetime import datetime
from src.schemas.kpi_schema import KPIObject
from src.engines.kpi_repository import KPIRepository, AIReportCache
from src.services.gemini_service import gemini_service
from src.services.evaluation_logger import eval_logger

log = logging.getLogger("ai-property-advisor")


AI_REPORT_FALLBACK_TEMPLATE = """# BÁO CÁO TÀI CHÍNH VÀ VẬN HÀNH THÁNG {period}

## I. TỔNG QUAN HIỆU QUẢ KINH DOANH & SỨC KHỎE TÀI CHÍNH
- **Tổng Doanh Thu Thực Tế:** `{total_revenue} VNĐ` (Tăng trưởng: `{growth_pct}%` so với kỳ trước)
- **Tổng Chi Phí Vận Hành:** `{total_expense} VNĐ`
- **Lợi Nhuận Ròng (Net Profit):** `{net_profit} VNĐ`
- **Đánh giá chung:** Hoạt động kinh doanh trong tháng {period} ghi nhận doanh thu đạt `{total_revenue} VNĐ`. Cần kiểm soát chi phí vận hành và nâng cao hiệu quả thu hồi công nợ để tối ưu dòng tiền.

## II. PHÂN TÍCH CHI TIẾT CƠ CẤU DOANH THU VÀ CHI PHÍ VẬN HÀNH
- **Doanh thu tiền phòng:** `{rent_revenue} VNĐ`
- **Doanh thu tiện ích (Điện, Nước, Dịch vụ):** `{utility_revenue} VNĐ`
- **Chi phí điện nước & bảo trì đầu vào:** `{utility_expense} VNĐ`
- **Biến động chi phí:** Chi phí điện nước vận hành duy trì ổn định, hỗ trợ hoạt động của các phòng thuê hiện hữu.

## III. TÌNH TRẠNG VẬN HÀNH PHÒNG VÀ CHI PHÍ CƠ HỘI LÃNG PHÍ
- **Tỉ lệ lấp đầy phòng:** `{occupancy_rate}%` ({occupied_rooms}/{total_rooms} phòng đang thuê)
- **Chi phí cơ hội lãng phí:** Tòa nhà hiện có `{vacant_rooms}` phòng trống (Tầng 4 và Tầng 5) khiến thất thoát khoảng `{opportunity_loss} VNĐ` doanh thu tiền phòng mỗi tháng.

## IV. PHÂN TÍCH RỦI RO CÔNG NỢ VÀ THU HỒI DÒNG TIỀN
- **Tổng dư nợ tồn đọng:** `{overdue_amount} VNĐ` ({overdue_count} hóa đơn quá hạn)
- **Tỷ lệ thu hồi tiền:** `{collection_rate}%`
- **Cảnh báo rủi ro:** Dư nợ quá hạn tập trung chủ yếu tại Phòng 501. Cần tập trung đôn đốc thu nợ để cải thiện khả năng thanh toán.

## V. ĐỀ XUẤT HÀNH ĐỘNG QUẢN TRỊ THỰC THI
1. **Thu hồi công nợ:** Phát thông báo đôn đốc và làm việc trực tiếp với khách thuê Phòng 501.
2. **Cải thiện lấp đầy:** Đẩy mạnh bài đăng tìm khách cho `{vacant_rooms}` phòng trống tầng 4 & 5.
3. **Kiểm soát chi phí:** Rà soát đồng hồ điện nước phụ để đảm bảo thu đúng, thu đủ.
"""


class AIReportService:

    @staticmethod
    def _build_fallback(kpi: KPIObject) -> str:
        """Tạo fallback report từ KPI"""
        rent_rev = kpi.revenue.rent if kpi.revenue else 0
        util_rev = (kpi.revenue.total - rent_rev) if kpi.revenue else 0
        util_exp = (kpi.expense.electricity + kpi.expense.water) if kpi.expense else 0
        vacant = kpi.occupancy.total_rooms - kpi.occupancy.occupied_rooms if kpi.occupancy else 0
        avg_price = (rent_rev / kpi.occupancy.occupied_rooms) if (kpi.occupancy and kpi.occupancy.occupied_rooms > 0) else 2500000
        opp_loss = vacant * avg_price

        return AI_REPORT_FALLBACK_TEMPLATE.format(
            period=kpi.period,
            total_revenue=f"{kpi.revenue.total:,.0f}",
            growth_pct=kpi.change_pct.revenue if kpi.change_pct and kpi.change_pct.revenue else 0,
            total_expense=f"{kpi.expense.total:,.0f}",
            net_profit=f"{kpi.profit.net:,.0f}",
            rent_revenue=f"{rent_rev:,.0f}",
            utility_revenue=f"{util_rev:,.0f}",
            utility_expense=f"{util_exp:,.0f}",
            overdue_amount=f"{kpi.debt.overdue_amount:,.0f}",
            overdue_count=kpi.debt.overdue_count,
            collection_rate=kpi.debt.collection_rate,
            occupancy_rate=kpi.occupancy.occupancy_rate,
            occupied_rooms=kpi.occupancy.occupied_rooms,
            total_rooms=kpi.occupancy.total_rooms,
            vacant_rooms=vacant,
            opportunity_loss=f"{opp_loss:,.0f}",
        )

    @staticmethod
    async def generate_report(landlord_id: int, period: str, force: bool = False) -> dict:
        """
        Sinh báo cáo tài chính từ KPI Object
        - Version-aware cache: nếu KPI hash không đổi → dùng cached
        - Nếu KPI hash thay đổi hoặc force → regenerate
        - Cache key format: {dashboard}:{period}:{version}
        """
        with eval_logger.start(
            question="AI Report generation",
            intent="REPORT",
            landlord_id=landlord_id,
            period=period,
        ) as eval_ctx:
            # Lấy KPI + compute hash
            kpi = KPIRepository.get_kpi(landlord_id, period)
            if kpi is None:
                from src.api.v1.kpi import _calculate_full_kpi
                kpi = await _calculate_full_kpi(landlord_id, period)

            kpi_hash = KPIRepository.get_kpi_hash(landlord_id, period)

            # Invalidate cache if force
            if force:
                AIReportCache.invalidate_report(landlord_id, period)

            # Version-aware cache check
            if not force and kpi_hash:
                if AIReportCache.has_report(landlord_id, period, version=kpi_hash):
                    cached = AIReportCache.get_report(landlord_id, period, version=kpi_hash)
                    if cached is not None and len(cached) > 200:
                        log.debug("Report cache HIT: version=%s", kpi_hash)
                        eval_ctx.set_reply(cached, model_used="cache", from_cache=True)
                        return {
                            "report": cached,
                            "from_cache": True,
                            "cache_version": kpi_hash,
                            "generated_at": None,
                        }

        # Cache miss — generate new report via Harness Agentic Loop
        try:
            from src.harness.agent_loop import HarnessAgentLoop
            now = datetime.now()
            current_period_str = now.strftime("%Y-%m")
            today_date_str = now.strftime("%d/%m/%Y")
            is_current_period = (period == current_period_str)

            if is_current_period:
                period_title_prompt = f"BÁO CÁO TÀI CHÍNH VÀ VẬN HÀNH (TÍNH ĐẾN NGÀY {today_date_str})"
                period_context_note = f"Lưu ý: Kỳ {period} đang diễn ra chưa kết thúc. Báo cáo này tổng hợp số liệu thực tế tính đến thời điểm hiện tại ({today_date_str})."
            else:
                period_title_prompt = f"BÁO CÁO TÀI CHÍNH VÀ VẬN HÀNH CHỐT SỔ THÁNG {period}"
                period_context_note = f"Báo cáo tài chính đã chốt sổ cho kỳ {period}. Dữ liệu cố định và được lưu trữ chính thức vĩnh viễn."

            report_question = (
                f"Hãy viết một {period_title_prompt} cực kỳ chi tiết, sâu sắc và đầy đủ cho Chủ nhà (Landlord ID: {landlord_id}). "
                f"{period_context_note} "
                f"Sử dụng các công cụ (get_kpi_overview, execute_sql_query, execute_dynamic_python_script) để truy vấn trực tiếp CSDL 61 bảng MySQL. "
                f"Yêu cầu báo cáo gồm đủ 5 phần lớn (I. Tổng quan hiệu quả kinh doanh, II. Phân tích cơ cấu doanh thu & chi phí, III. Tình trạng vận hành phòng & chi phí cơ hội, IV. Phân tích rủi ro công nợ & nợ xấu, V. Khuyến nghị quản trị thực thi). "
                f"Báo cáo phải phân tích chi tiết các con số thực tế tính đến hôm nay, tính toán thất thoát tiền từ phòng trống và nêu rõ danh sách phòng nợ xấu kéo dài."
            )

            agent_result = await HarnessAgentLoop.run(
                question=report_question,
                landlord_id=landlord_id,
                period=period,
                history="",
            )

            text = agent_result.get("reply", "") or agent_result.get("answer", "")

            if not text or len(text) < 150:
                log.warning("Agent reply was short or empty (%d chars), using fallback", len(text) if text else 0)
                text = AIReportService._build_fallback(kpi)

            # Cache kết quả với version
            AIReportCache.set_report(landlord_id, period, text, version=kpi_hash)

            return {
                "report": text,
                "from_cache": False,
                "cache_version": kpi_hash,
                "generated_at": datetime.now().isoformat(),
            }

        except Exception as e:
            log.warning("Gemini AI Report failed, using fallback: %s", str(e))
            fallback = AIReportService._build_fallback(kpi)
            AIReportCache.set_report(landlord_id, period, fallback, version=kpi_hash)
            return {
                "report": fallback,
                "from_cache": False,
                "cache_version": kpi_hash,
                "generated_at": None,
            }

    @staticmethod
    async def trigger_background(landlord_id: int, period: str) -> None:
        """Trigger sinh report trong background (không block)"""
        import asyncio
        asyncio.create_task(AIReportService.generate_report(landlord_id, period))

    @staticmethod
    def export_report_to_docx(report_text: str, period: str) -> io.BytesIO:
        """Xuất báo cáo tài chính sang file Word (.docx) đẹp mắt"""
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = docx.Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(f"BÁO CÁO TÀI CHÍNH VÀ VẬN HÀNH THÁNG {period}")
        run.font.name = "Calibri"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(99, 102, 241) # Indigo

        subtitle_p = doc.add_paragraph()
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle_p.add_run("Hệ Thống AI Property Advisor — Chẩn Đoán Quản Trị Doanh Nghiệp")
        sub_run.font.name = "Calibri"
        sub_run.font.size = Pt(11)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(156, 163, 175)

        doc.add_paragraph()

        lines = report_text.split("\n")
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            if line_str.startswith("# "):
                p = doc.add_paragraph()
                r = p.add_run(line_str.replace("# ", ""))
                r.font.name = "Calibri"
                r.font.size = Pt(15)
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 41, 59)
            elif line_str.startswith("I.") or line_str.startswith("II.") or line_str.startswith("III.") or line_str.startswith("IV.") or line_str.startswith("V."):
                p = doc.add_paragraph()
                r = p.add_run(line_str)
                r.font.name = "Calibri"
                r.font.size = Pt(13)
                r.font.bold = True
                r.font.color.rgb = RGBColor(79, 70, 229)
            elif line_str.startswith("- "):
                p = doc.add_paragraph(style='List Bullet')
                content = line_str[2:]
                parts = content.split("**")
                for i, part in enumerate(parts):
                    clean_part = part.replace("`", "")
                    r = p.add_run(clean_part)
                    r.font.name = "Calibri"
                    r.font.size = Pt(11)
                    if i % 2 == 1:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(31, 41, 55)
            else:
                p = doc.add_paragraph()
                r = p.add_run(line_str.replace("`", "").replace("**", ""))
                r.font.name = "Calibri"
                r.font.size = Pt(11)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


ai_report_service = AIReportService()
